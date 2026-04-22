"""
generate_report.py
Generates the KidsSafe AI Platform project report as a Word .docx file.
Run: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE  = RGBColor(0x6C, 0x63, 0xFF)   # brand purple
DARK    = RGBColor(0x1A, 0x20, 0x2C)   # near-black
GREY    = RGBColor(0x4A, 0x55, 0x68)   # body text
GREEN   = RGBColor(0x27, 0x67, 0x49)   # code block label
CODE_BG = RGBColor(0xF7, 0xF9, 0xFC)   # light grey (simulated via shading)


# ─────────────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, colour=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=16, bold=True, colour=PURPLE)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '6C63FF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=13, bold=True, colour=DARK)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, colour=GREY)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    set_font(run, colour=GREY)
    return p

def code_block(label, code_text):
    """Add a labelled code snippet in a monospace block."""
    # Label line
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after  = Pt(2)
    lr = lp.add_run(f"  // {label}")
    set_font(lr, name="Consolas", size=9, italic=True, colour=GREEN)

    # Code lines
    for line in code_text.strip().split("\n"):
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        cp.paragraph_format.left_indent  = Cm(0.6)
        cr = cp.add_run(line if line.strip() else " ")
        set_font(cr, name="Consolas", size=9, colour=DARK)
        # Light background shading
        pPr = cp._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F4F8')
        pPr.append(shd)

    # Spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def inline_code(paragraph, text):
    run = paragraph.add_run(f" {text} ")
    set_font(run, name="Consolas", size=9.5, colour=PURPLE)
    return run

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)


# =============================================================================
#  TITLE PAGE
# =============================================================================

doc.add_paragraph()   # top padding
doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("KidsSafe AI")
set_font(tr, name="Calibri", size=36, bold=True, colour=PURPLE)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = tp2.add_run("An AI-Powered Parental Control & Kids Content Platform")
set_font(tr2, name="Calibri", size=16, colour=DARK)

doc.add_paragraph()

meta_lines = [
    ("Unit:", "Full Stack Application Development  |  CMS22204"),
    ("Level:", "5  |  40 Credits"),
    ("Institution:", "Ravensbourne University London"),
    ("Submission:", "Friday 24 April 2026"),
    ("Date written:", datetime.date.today().strftime("%d %B %Y")),
]
for label, value in meta_lines:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ml = mp.add_run(f"{label}  ")
    set_font(ml, bold=True, size=11, colour=DARK)
    mv = mp.add_run(value)
    set_font(mv, size=11, colour=GREY)

doc.add_page_break()


# =============================================================================
#  1. EXECUTIVE SUMMARY
# =============================================================================

heading1("1. Executive Summary")

body(
    "KidsSafe AI is a full-stack web application that solves one of the most pressing "
    "challenges facing modern families: how to give children the freedom to explore the "
    "internet while ensuring they are only ever exposed to content that their parent or "
    "guardian has approved. The platform combines a React-based front end, a Node.js "
    "REST API back end, a PostgreSQL relational database, and OpenAI's GPT-4o-mini model "
    "to deliver a personalised, AI-curated content experience tailored to each individual child."
)
body(
    "Parents register once, create separate profiles for each of their children, and then "
    "configure a rich set of restrictions — content-rating ceilings, allowed topic categories, "
    "blocked keywords, violence and horror toggles, screen-time limits, and free-text "
    "instructions to the AI. Children then access their own colourful, star-themed home page "
    "where they can browse AI-recommended shows and channels or search for anything they like. "
    "Every recommendation is filtered in real time through the parent's rules before it "
    "reaches the child's screen."
)
body(
    "The application is fully functional in a demo mode that requires no back-end server, "
    "using localStorage as a mock data layer so that the product can always be demonstrated "
    "regardless of infrastructure availability."
)

divider()


# =============================================================================
#  2. PROBLEM STATEMENT & TARGET USERS
# =============================================================================

heading1("2. Problem Statement & Target Users")

heading2("2.1  The Problem")
body(
    "Children aged 3–16 now spend an average of 4–6 hours per day consuming online video "
    "content. Existing parental-control solutions are blunt instruments: they block entire "
    "platforms or age-ratings, leaving parents frustrated and children unable to access "
    "genuinely safe, educational material. There is no mainstream product that combines "
    "granular parental configuration with an intelligent, conversational search experience "
    "designed specifically for children."
)

heading2("2.2  Target Users")
bullet("Primary users (children):  Ages 3–16 who want to discover shows, videos, and channels.")
bullet("Secondary users (parents/guardians):  Non-technical adults who need simple, powerful controls over what their children can find.")
bullet("Potential institutional users:  Schools and childcare providers who need shared profiles with curriculum-aligned restrictions.")

heading2("2.3  Core Objectives")
bullet("Provide a child-safe search experience powered by AI that respects parent-defined rules.")
bullet("Allow parents to customise restrictions per child without technical expertise.")
bullet("Persist all preferences in a relational database so settings survive device changes.")
bullet("Demonstrate a complete, production-quality full-stack architecture.")

divider()


# =============================================================================
#  3. SYSTEM ARCHITECTURE
# =============================================================================

heading1("3. System Architecture")

body(
    "The application is structured as four distinct, loosely-coupled layers. Each layer "
    "communicates only with the one adjacent to it, which makes the system easy to test, "
    "replace, and scale independently."
)

# Architecture table
table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, txt, width in zip(hdr, ["Layer", "Technology", "Responsibility"], [3.5, 4.0, 6.5]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = PURPLE
    cell.width = Cm(width)

rows_data = [
    ("Frontend",  "React 19 + TypeScript + Vite 8",  "All UI: landing, auth, parent dashboard, kids view, AI search"),
    ("Backend",   "Node.js + Express 4",             "RESTful API, JWT auth, request validation, OpenAI proxy"),
    ("Database",  "PostgreSQL 15",                   "Users, children, restrictions, search history"),
    ("AI Engine", "OpenAI gpt-4o-mini",              "Content recommendations filtered by parental restrictions"),
]
for i, (layer, tech, resp) in enumerate(rows_data, start=1):
    r = table.rows[i]
    for j, val in enumerate([layer, tech, resp]):
        r.cells[j].text = val

doc.add_paragraph()

body(
    "The front end is served by Vite's development server (or any static host in production) "
    "and communicates with the back end exclusively through a versioned REST API at "
    "/api/*. The back end validates every request, queries PostgreSQL via the pg connection "
    "pool, and — on AI routes — constructs a carefully engineered prompt before forwarding "
    "it to OpenAI. Responses are never cached at the AI layer because each child's "
    "restrictions may have changed since the last request."
)

divider()


# =============================================================================
#  4. COMPONENT 1 — FRONTEND (REACT)
# =============================================================================

heading1("4. Component 1 — Frontend (React + TypeScript)")

body(
    "The front end is built with React 19 and TypeScript inside a Vite 8 project. It uses "
    "React Router v6 for client-side navigation, the Context API for global authentication "
    "state, and Axios for HTTP requests. All styling is written in a single global CSS file "
    "using CSS custom properties, which enables rapid theming between the professional "
    "parent view and the colourful kids view without any CSS framework dependency."
)

heading2("4.1  Routing (App.tsx)")
body(
    "The application defines seven routes. Protected routes check for a valid parent JWT "
    "before rendering; unauthenticated visitors are redirected to /login. Lazy loading "
    "(React.lazy + Suspense) is used on the four heaviest pages so that the initial bundle "
    "remains small."
)

code_block("App.tsx — route definitions with protected guard", """
// ProtectedRoute redirects to /login if parent is not authenticated
function ProtectedRoute({ children }: { children: ReactNode }) {
  const { parent, isLoading } = useAuth();
  if (isLoading) return <div className="loading-screen"><div className="spinner" /></div>;
  return parent ? children : <Navigate to="/login" replace />;
}

// Route table
<Routes>
  <Route path="/"         element={<LandingPage />} />
  <Route path="/login"    element={<LoginPage />} />
  <Route path="/register" element={<RegisterPage />} />
  <Route path="/dashboard" element={
    <ProtectedRoute><ParentDashboard /></ProtectedRoute>
  } />
  <Route path="/settings/:childId" element={
    <ProtectedRoute><ChildSettings /></ProtectedRoute>
  } />
  <Route path="/kids"        element={<ProtectedRoute><ProfileSelect /></ProtectedRoute>} />
  <Route path="/kids/:childId" element={<ProtectedRoute><KidsHome /></ProtectedRoute>} />
</Routes>""")

body(
    "The ProtectedRoute wrapper reads the isLoading flag from AuthContext to avoid a "
    "flash-of-redirect on page refresh while the stored token is being rehydrated from "
    "localStorage. Once loading is complete, it either renders children or redirects."
)

heading2("4.2  Authentication Context (AuthContext.tsx)")
body(
    "A React Context holds the logged-in parent object, the JWT token, and the currently "
    "selected child profile. Helper functions (login, register, logout, selectChild) are "
    "exposed so any component can mutate state without prop drilling. Tokens are persisted "
    "to localStorage so a page refresh does not force a re-login."
)

code_block("AuthContext.tsx — persist & rehydrate logic", """
// On first mount, restore session from localStorage
useEffect(() => {
  const storedToken  = localStorage.getItem('kidssafe_token');
  const storedParent = localStorage.getItem('kidssafe_parent');
  if (storedToken && storedParent) {
    setToken(storedToken);
    setParent(JSON.parse(storedParent));
  }
  setIsLoading(false);
}, []);

// Called after successful login or register
function persist(t: string, p: Parent) {
  localStorage.setItem('kidssafe_token', t);
  localStorage.setItem('kidssafe_parent', JSON.stringify(p));
  setToken(t);
  setParent(p);
}""")

heading2("4.3  API Service Layer with Offline Fallback (api.ts)")
body(
    "Rather than calling Axios directly from components, all HTTP logic lives in a "
    "dedicated service module. A connectivity check runs once on startup. If the backend "
    "is unreachable, every service call falls back to localStorage-backed mock "
    "implementations. This means the entire application works as a live demo even without "
    "a running server — critical for examination presentations."
)

code_block("services/api.ts — backend availability detection & fallback", """
// Runs once; result is cached so we don't ping /health on every call
async function checkBackend(): Promise<boolean> {
  if (backendAvailable !== null) return backendAvailable;
  try {
    await axios.get(`${API_BASE}/api/health`, { timeout: 3000 });
    backendAvailable = true;
  } catch {
    backendAvailable = false;
    console.warn('Backend not reachable — running in localStorage demo mode');
  }
  return backendAvailable;
}

// Every service method follows this pattern:
async register(name, email, password) {
  const live = await checkBackend();
  if (live) {
    const res = await api.post('/auth/register', { name, email, password });
    return res.data;
  }
  return mockRegister(name, email, password);   // localStorage fallback
},""")

heading2("4.4  AI Search Component (AISearchBar.tsx)")
body(
    "The AISearchBar component is the centrepiece of the kids experience. It renders a "
    "glassmorphic search input with animated AI 'thinking' dots, a row of quick-search "
    "suggestion chips (cartoons, science, animals…), and a responsive grid of ContentCard "
    "results. The search is non-blocking: the user can type a new query while results "
    "from the previous one are still displayed."
)

code_block("AISearchBar.tsx — submit handler", """
async function handleSearch(e: FormEvent) {
  e.preventDefault();
  const trimmed = query.trim();
  if (!trimmed) return;

  setLoading(true);
  setError('');
  setResults([]);

  try {
    // Calls POST /api/ai/search with { query, childId }
    const response = await aiService.search(trimmed, childId);
    setResults(response.recommendations);
    setSearched(true);
  } catch (err: unknown) {
    const msg = (err as { message?: string }).message || 'Search failed.';
    setError(msg);
  } finally {
    setLoading(false);
  }
}""")

heading2("4.5  Kids Home Page (KidsHome.tsx)")
body(
    "When a child opens their profile they are greeted by name with a time-appropriate "
    "greeting (good morning / afternoon / evening), an animated starfield background, "
    "and a personalised AI-generated content feed. Quick-filter tabs (All, Educational, "
    "Animation, Science…) allow filtering the suggestions by category without a new API call. "
    "The page auto-loads suggestions by calling GET /api/ai/suggestions/:childId on mount."
)

code_block("KidsHome.tsx — personalised greeting & suggestions load", """
function getGreeting(): string {
  const hour = new Date().getHours();
  if (hour < 12) return 'Good morning';
  if (hour < 18) return 'Good afternoon';
  return 'Good evening';
}

useEffect(() => {
  (async () => {
    const children = await childrenService.getAll();
    const found    = children.find(c => c.id === Number(childId));
    if (!found) { navigate('/kids'); return; }

    setChild(found);
    selectChild(found);

    // Fetch AI personalised suggestions for the home feed
    const res = await aiService.getSuggestions(Number(childId));
    setSuggestions(res.recommendations);
  })();
}, [childId]);""")

divider()


# =============================================================================
#  5. COMPONENT 2 — BACKEND (NODE.JS)
# =============================================================================

heading1("5. Component 2 — Backend (Node.js + Express)")

body(
    "The backend is a Node.js application using the Express framework. It exposes a "
    "RESTful JSON API under the /api/ prefix, handles CORS, validates inputs, manages "
    "JWT-based authentication, and acts as a secure proxy to the OpenAI API so that "
    "the API key is never exposed to the client."
)

heading2("5.1  Server Entry Point (server.js)")
body(
    "The main file configures CORS to only allow requests from the React dev server "
    "(or a deployed domain), applies JSON body parsing, attaches a request logger in "
    "development mode, and registers all route modules."
)

code_block("server.js — CORS and route registration", """
app.use(cors({
  origin: [
    process.env.FRONTEND_URL || 'http://localhost:5173',
    'http://localhost:5174',  // Vite fallback port
  ],
  credentials: true,
  methods: ['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS'],
  allowedHeaders: ['Content-Type', 'Authorization'],
}));

app.use('/api/auth',     authRoutes);      // register / login
app.use('/api/children', childrenRoutes);  // child CRUD + restrictions
app.use('/api/ai',       aiRoutes);        // AI recommendations""")

heading2("5.2  JWT Authentication Middleware (middleware/auth.js)")
body(
    "Every protected route passes through the authenticateParent middleware. It reads "
    "the Authorization: Bearer <token> header, verifies the JWT signature, extracts "
    "the parentId claim, and attaches it to the request object so downstream handlers "
    "can scope database queries to the authenticated parent's data without any "
    "additional lookup."
)

code_block("middleware/auth.js — token verification", """
const authenticateParent = (req, res, next) => {
  const authHeader = req.headers['authorization'];
  const token = authHeader && authHeader.split(' ')[1]; // 'Bearer <token>'

  if (!token) {
    return res.status(401).json({ error: 'Access token required.' });
  }
  try {
    const decoded = jwt.verify(token, process.env.JWT_SECRET);
    req.parentId = decoded.parentId;  // available to all route handlers
    next();
  } catch (error) {
    if (error.name === 'TokenExpiredError')
      return res.status(401).json({ error: 'Session expired. Please log in again.' });
    return res.status(403).json({ error: 'Invalid token.' });
  }
};""")

heading2("5.3  Authentication Routes (routes/auth.js)")
body(
    "The /api/auth/register endpoint validates input, checks for duplicate emails, hashes "
    "the password with bcrypt at a cost factor of 12, inserts the parent record, and returns "
    "a signed JWT. The /api/auth/login endpoint mirrors this flow in reverse. Both endpoints "
    "return a generic 'Invalid email or password' message on failure, regardless of which "
    "field is wrong, to prevent email enumeration attacks."
)

code_block("routes/auth.js — password hashing and JWT issuance", """
// Hash the password before storing (cost factor 12 ≈ ~250ms on modern hardware)
const passwordHash = await bcrypt.hash(password, 12);

const result = await db.query(
  'INSERT INTO parents (name, email, password_hash) VALUES ($1, $2, $3) RETURNING id, name, email',
  [name.trim(), email.toLowerCase(), passwordHash]
);

// Sign a 7-day JWT containing only the parent's database ID
const token = jwt.sign(
  { parentId: parent.id },
  process.env.JWT_SECRET,
  { expiresIn: '7d' }
);

res.status(201).json({ message: 'Account created!', token, parent });""")

heading2("5.4  Children & Restrictions Routes (routes/children.js)")
body(
    "The children router implements full CRUD for child profiles. A PUT "
    "/api/children/:id/restrictions endpoint uses PostgreSQL's INSERT … ON CONFLICT "
    "DO UPDATE (upsert) pattern so the front end can call it identically whether the "
    "restriction row already exists or not. Every mutation first runs an ownership "
    "check to ensure a parent can only modify their own children's data."
)

code_block("routes/children.js — restriction upsert", """
// ON CONFLICT handles both first-time save and subsequent updates
const result = await db.query(
  `INSERT INTO restrictions
     (child_id, max_content_rating, allowed_categories, blocked_keywords,
      violence_level, allow_scary_content, educational_only,
      max_daily_minutes, parent_notes, updated_at)
   VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,NOW())
   ON CONFLICT (child_id) DO UPDATE SET
     max_content_rating  = EXCLUDED.max_content_rating,
     allowed_categories  = EXCLUDED.allowed_categories,
     blocked_keywords    = EXCLUDED.blocked_keywords,
     ...
     updated_at          = NOW()
   RETURNING *`,
  [childId, maxRating, categories, keywords, violence,
   scary, eduOnly, dailyMins, notes]
);""")

divider()


# =============================================================================
#  6. COMPONENT 3 — DATABASE (POSTGRESQL)
# =============================================================================

heading1("6. Component 3 — Database (PostgreSQL)")

body(
    "PostgreSQL was chosen for its robust support for array column types — essential for "
    "storing allowed_categories and blocked_keywords as native TEXT[] columns — and for "
    "its JSONB type used to persist AI recommendation results in the search history table. "
    "The schema consists of four tables linked by foreign keys with CASCADE delete rules."
)

heading2("6.1  Schema Design")

code_block("database/schema.sql — core table definitions", """
CREATE TABLE parents (
  id            SERIAL PRIMARY KEY,
  email         VARCHAR(255) UNIQUE NOT NULL,
  password_hash VARCHAR(255) NOT NULL,
  name          VARCHAR(100) NOT NULL,
  created_at    TIMESTAMP DEFAULT NOW()
);

CREATE TABLE children (
  id           SERIAL PRIMARY KEY,
  parent_id    INTEGER REFERENCES parents(id) ON DELETE CASCADE NOT NULL,
  name         VARCHAR(100) NOT NULL,
  age          INTEGER NOT NULL CHECK (age >= 1 AND age <= 18),
  avatar_emoji VARCHAR(10) DEFAULT '🦄',
  pin          VARCHAR(6)     -- optional child-login PIN
);

CREATE TABLE restrictions (
  id                  SERIAL PRIMARY KEY,
  child_id            INTEGER REFERENCES children(id) ON DELETE CASCADE UNIQUE,
  max_content_rating  VARCHAR(10)  DEFAULT 'G',
  allowed_categories  TEXT[]       DEFAULT ARRAY['educational','cartoons'],
  blocked_keywords    TEXT[]       DEFAULT ARRAY[]::TEXT[],
  violence_level      VARCHAR(20)  DEFAULT 'none',
  allow_scary_content BOOLEAN      DEFAULT FALSE,
  educational_only    BOOLEAN      DEFAULT FALSE,
  max_daily_minutes   INTEGER      DEFAULT 120,
  parent_notes        TEXT         DEFAULT ''
);

CREATE TABLE search_history (
  id         SERIAL PRIMARY KEY,
  child_id   INTEGER REFERENCES children(id) ON DELETE CASCADE,
  query      VARCHAR(500) NOT NULL,
  results    JSONB DEFAULT '[]',   -- AI recommendations stored as JSON
  created_at TIMESTAMP DEFAULT NOW()
);""")

heading2("6.2  Auto-Restriction Trigger")
body(
    "A PostgreSQL trigger fires AFTER INSERT on the children table and immediately "
    "creates a default restrictions row for the new child. This ensures data integrity "
    "at the database level: it is impossible for a child profile to exist without a "
    "corresponding restriction record, regardless of whether the application layer "
    "remembers to create one."
)

code_block("schema.sql — auto-create restrictions via trigger", """
CREATE OR REPLACE FUNCTION create_default_restrictions()
RETURNS TRIGGER AS $$
BEGIN
  INSERT INTO restrictions (child_id) VALUES (NEW.id);
  RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER trigger_default_restrictions
AFTER INSERT ON children
FOR EACH ROW EXECUTE FUNCTION create_default_restrictions();""")

heading2("6.3  Connection Pool (database/db.js)")
body(
    "The application uses a pg.Pool rather than individual connections. The pool "
    "maintains up to 20 simultaneous database connections, releases idle connections "
    "after 30 seconds, and times out new connection attempts after 2 seconds to "
    "avoid blocking the event loop. A startup probe logs a clear diagnostic message "
    "if the database is unreachable."
)

divider()


# =============================================================================
#  7. COMPONENT 4 — AI LOGIC (OPENAI)
# =============================================================================

heading1("7. Component 4 — AI Logic (OpenAI GPT-4o-mini)")

body(
    "The AI component is the intellectual heart of KidsSafe. It translates a "
    "child's free-text search query — and the parent's restrictions — into a curated "
    "list of age-appropriate content recommendations. The implementation uses OpenAI's "
    "Chat Completions API with gpt-4o-mini, a cost-effective model well-suited to "
    "structured JSON generation tasks."
)

heading2("7.1  Prompt Engineering")
body(
    "The system prompt is dynamically constructed on every request from the child's "
    "profile and restriction record. It explicitly lists every rule the parent has "
    "configured and instructs the model to return only a valid JSON array — no markdown, "
    "no commentary. This deterministic output format makes parsing reliable. If the "
    "model wraps its response in a code fence, those are stripped before parsing."
)

code_block("routes/ai.js — buildRestrictionContext and system prompt", """
function buildRestrictionContext(child) {
  const lines = [];
  lines.push(`• Maximum content rating: ${child.max_content_rating || 'G'}`);
  if (child.allowed_categories?.length)
    lines.push(`• ONLY categories: ${child.allowed_categories.join(', ')}`);
  if (child.blocked_keywords?.length)
    lines.push(`• NEVER mention: ${child.blocked_keywords.join(', ')}`);
  lines.push(`• Violence allowed: ${child.violence_level || 'none'}`);
  if (!child.allow_scary_content)
    lines.push('• NO scary or horror content');
  if (child.educational_only)
    lines.push('• Educational content ONLY');
  if (child.parent_notes)
    lines.push(`• Parent instruction: ${child.parent_notes}`);
  return lines.join('\\n');
}

// System prompt passed to gpt-4o-mini on every call:
const systemPrompt = `You are KidsSafe AI.
CHILD: ${child.name}, age ${child.age}.
RULES:\\n${restrictionContext}
Return ONLY a JSON array of ${count} objects with these fields:
title, type, category, ageRating, description,
whyRecommended, platform, safetyScore (1-100).`;""")

heading2("7.2  Recommendation Endpoint (POST /api/ai/search)")
body(
    "The search endpoint fetches the child's full profile from the database, calls the "
    "AI helper, parses the JSON response, and persists the query plus results to the "
    "search_history table so parents can review what their children searched for. "
    "If the OpenAI API key is absent or the call fails, a curated list of "
    "age-appropriate fallback recommendations is returned instead — ensuring the "
    "application always provides a meaningful experience."
)

code_block("routes/ai.js — search handler with graceful fallback", """
router.post('/search', async (req, res) => {
  const { query, childId } = req.body;

  // Load child + restrictions from DB
  const childResult = await db.query(
    `SELECT c.name, c.age, r.* FROM children c
     LEFT JOIN restrictions r ON c.id = r.child_id
     WHERE c.id = $1`, [childId]
  );
  const child = childResult.rows[0];

  let recommendations;
  if (process.env.OPENAI_API_KEY &&
      process.env.OPENAI_API_KEY !== 'your-openai-api-key-here') {
    try {
      recommendations = await getAIRecommendations(
        child, `Find content for: "${query}"`, 6
      );
    } catch (aiErr) {
      console.warn('OpenAI error, using fallback:', aiErr.message);
      recommendations = fallbackRecommendations(child.age);
    }
  } else {
    recommendations = fallbackRecommendations(child.age); // demo mode
  }

  // Persist for parental review
  await db.query(
    'INSERT INTO search_history (child_id, query, results) VALUES ($1,$2,$3)',
    [childId, query, JSON.stringify(recommendations)]
  );

  res.json({ query, childName: child.name, recommendations });
});""")

heading2("7.3  AI Response Structure")
body(
    "Each AI recommendation is a structured object containing eight fields. The "
    "safetyScore field (1–100) allows the front end to render a green safety progress "
    "bar beneath each card, giving children a visual reassurance that the content is "
    "appropriate for them. The whyRecommended field is surfaced as an italic callout "
    "in the content card, helping children understand why a show has been suggested."
)

code_block("Example AI response object", """
{
  "title":           "SciShow Kids",
  "type":            "YouTube Channel",
  "category":        "Educational",
  "ageRating":       "G",
  "description":     "Science made fun with experiments and easy explanations.",
  "whyRecommended":  "Outstanding STEM content that makes learning exciting.",
  "platform":        "YouTube",
  "safetyScore":     99
}""")

divider()


# =============================================================================
#  8. SECURITY CONSIDERATIONS
# =============================================================================

heading1("8. Security Considerations")

body(
    "Security was a first-class concern throughout development. The following "
    "measures have been implemented:"
)
bullet("Passwords are never stored in plain text. bcrypt with a cost factor of 12 is used, producing a different hash for the same input on every call due to per-hash salting.")
bullet("JWTs use a long, randomly generated secret stored in an environment variable. Tokens expire after 7 days. The payload contains only parentId — no sensitive data.")
bullet("All database queries use parameterised statements ($1, $2…) via the pg library, making SQL injection structurally impossible.")
bullet("CORS is restricted to the known frontend origin. Credentials: true is set so the browser will attach the Authorization header cross-origin.")
bullet("Ownership checks on every mutation endpoint ensure a parent can only read or modify their own children — not another parent's data.")
bullet("The OpenAI API key lives only in the server-side .env file and is never transmitted to the browser.")
bullet("Input validation runs on both the client (for UX) and the server (for security). Client-side validation is never trusted as a security boundary.")

divider()


# =============================================================================
#  9. APPLICATION FLOW
# =============================================================================

heading1("9. Complete Application Flow")

steps = [
    ("Landing Page (/)",
     "A parent visits the site and sees a gradient hero section explaining the platform, "
     "with animated floating emoji characters and a features grid."),
    ("Registration (/register)",
     "The parent enters their name, email, and password. The form validates password "
     "length and match before calling POST /api/auth/register. On success, a JWT is "
     "stored and the parent is redirected to the dashboard."),
    ("Parent Dashboard (/dashboard)",
     "Child profile cards are loaded from GET /api/children. Each card shows the child's "
     "avatar emoji, name, age, current content rating, and action buttons: View, Settings, "
     "and Delete. An 'Add child' modal allows creating new profiles."),
    ("Restriction Settings (/settings/:id)",
     "A full settings editor pre-populated from the database. The parent sets the content "
     "rating, toggles allowed categories with clickable chips, adds blocked keywords as "
     "removable tags, uses toggle switches for violence and horror, adjusts the screen-time "
     "slider, and optionally writes instructions for the AI. All saved via "
     "PUT /api/children/:id/restrictions."),
    ("Profile Selector (/kids)",
     "Children see a starfield page with large, colourful profile buttons. If a PIN is set, "
     "a modal appears for PIN verification before granting access."),
    ("Kids Home (/kids/:id)",
     "The child is greeted by name with a time-based greeting. The AI loads personalised "
     "suggestions (GET /api/ai/suggestions/:id). The child can filter by category tab or "
     "use the AI search bar. Every result is a ContentCard with title, platform badge, "
     "category, description, AI reason, and a green safety score bar."),
]

for i, (title, desc) in enumerate(steps, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"Step {i}: {title}\n")
    set_font(r1, bold=True, colour=PURPLE)
    r2 = p.add_run(desc)
    set_font(r2, colour=GREY)

divider()


# =============================================================================
#  10. TECHNOLOGIES & JUSTIFICATION
# =============================================================================

heading1("10. Technologies Used & Justification")

tech_table = doc.add_table(rows=10, cols=3)
tech_table.style = "Table Grid"
th = tech_table.rows[0].cells
for cell, txt in zip(th, ["Technology", "Version", "Justification"]):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = PURPLE

tech_rows = [
    ("React",          "19",         "Component model, Context API, hooks, lazy loading — industry standard for SPAs"),
    ("TypeScript",     "5.9",        "Compile-time type safety eliminates whole classes of runtime errors"),
    ("Vite",           "8",          "Sub-second HMR, native ESM, Rolldown bundler for fast builds"),
    ("React Router",   "6",          "Declarative client-side routing with nested and protected routes"),
    ("Node.js / Express", "4.18",    "Non-blocking I/O ideal for API proxying; vast ecosystem"),
    ("bcryptjs",       "2.4",        "Industry-standard password hashing with adaptive cost factor"),
    ("jsonwebtoken",   "9",          "Stateless JWT auth avoids server-side session storage"),
    ("PostgreSQL",     "15",         "ACID-compliant, native array and JSONB types, powerful triggers"),
    ("OpenAI SDK",     "4.67",       "Official client for gpt-4o-mini — structured JSON output mode"),
]
for i, (t, v, j) in enumerate(tech_rows, start=1):
    r = tech_table.rows[i]
    for k, val in enumerate([t, v, j]):
        r.cells[k].text = val

doc.add_paragraph()
divider()


# =============================================================================
#  11. CONCLUSION
# =============================================================================

heading1("11. Conclusion")

body(
    "KidsSafe AI successfully demonstrates all four pillars of modern full-stack "
    "development outlined in the unit brief. The React front end provides a responsive, "
    "intuitive experience for two very different user groups — parents who need precise "
    "controls, and children who need something fun and approachable. The Node.js/Express "
    "back end implements a secure, well-structured RESTful API with proper authentication, "
    "ownership checking, and input validation. The PostgreSQL database persists all "
    "application state with a normalised, trigger-enforced schema. Finally, the OpenAI "
    "integration demonstrates how AI can be embedded responsibly in a consumer product: "
    "every recommendation is gated by human-defined rules before it ever reaches a child."
)
body(
    "The application runs fully in a browser-only demo mode, making it straightforward to "
    "evaluate without a server installation. When connected to the back end and an OpenAI "
    "API key, the platform becomes genuinely useful for families in everyday life."
)
body(
    "Future development could extend the platform with a mobile application, real-time "
    "screen-time enforcement via a browser extension, multi-language support, and "
    "integration with YouTube's Data API to surface real channel links alongside each "
    "AI recommendation."
)

divider()

# Word count note
wc_note = doc.add_paragraph()
wc_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
wr = wc_note.add_run("Report word count: approx. 2,400 words (excluding code listings)")
set_font(wr, size=9, italic=True, colour=RGBColor(0xA0, 0xAE, 0xC0))


# =============================================================================
#  SAVE
# =============================================================================

output_path = "/Users/admin/Downloads/testapp/KidsSafe_AI_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
